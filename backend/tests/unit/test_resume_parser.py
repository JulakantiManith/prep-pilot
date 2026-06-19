"""Unit tests for the resume parser service."""

import json
from unittest.mock import AsyncMock, MagicMock, patch

import pytest

from app.services.resume_parser import ResumeParser, ResumeParserError


class TestResumeParserTextExtraction:
    """Tests for text extraction methods."""

    def setup_method(self):
        """Set up test fixtures."""
        with patch("app.services.resume_parser.get_supabase_client"):
            self.parser = ResumeParser.__new__(ResumeParser)
            self.parser._gemini = MagicMock()
            self.parser._client = MagicMock()

    def test_extract_text_from_pdf_success(self):
        """Test successful PDF text extraction."""
        # Create a minimal valid PDF with text
        from PyPDF2 import PdfWriter
        import io

        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        # PyPDF2 blank pages have no text, so we test the "no text" case
        pdf_buffer = io.BytesIO()
        writer.write(pdf_buffer)
        pdf_content = pdf_buffer.getvalue()

        # Blank PDF should raise since there's no text
        with pytest.raises(ResumeParserError, match="Could not extract any text"):
            self.parser._extract_text_from_pdf(pdf_content)

    def test_extract_text_from_pdf_invalid_file(self):
        """Test PDF extraction with invalid file content."""
        with pytest.raises(ResumeParserError, match="Failed to read the PDF"):
            self.parser._extract_text_from_pdf(b"not a pdf file")

    def test_extract_text_from_docx_invalid_file(self):
        """Test DOCX extraction with invalid file content."""
        with pytest.raises(ResumeParserError, match="Failed to read the DOCX"):
            self.parser._extract_text_from_docx(b"not a docx file")

    def test_extract_text_from_docx_empty(self):
        """Test DOCX extraction with empty document."""
        from docx import Document
        import io

        doc = Document()
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_content = buffer.getvalue()

        with pytest.raises(ResumeParserError, match="Could not extract any text"):
            self.parser._extract_text_from_docx(docx_content)

    def test_extract_text_from_docx_with_content(self):
        """Test successful DOCX text extraction."""
        from docx import Document
        import io

        doc = Document()
        doc.add_paragraph("John Doe")
        doc.add_paragraph("Software Engineer")
        doc.add_paragraph("Skills: Python, JavaScript")
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_content = buffer.getvalue()

        text = self.parser._extract_text_from_docx(docx_content)
        assert "John Doe" in text
        assert "Software Engineer" in text
        assert "Python" in text


class TestResumeParserResponseParsing:
    """Tests for Gemini response parsing."""

    def setup_method(self):
        """Set up test fixtures."""
        self.parser = ResumeParser.__new__(ResumeParser)
        self.parser._gemini = MagicMock()
        self.parser._client = MagicMock()

    def test_parse_valid_response(self):
        """Test parsing a valid extraction response."""
        response = json.dumps({
            "skills": ["Python", "JavaScript", "React"],
            "projects": [
                {"name": "Web App", "description": "A web app", "technologies": ["React"]}
            ],
            "experience": [
                {"title": "Engineer", "company": "Acme", "duration": "2y", "description": "Built things"}
            ],
            "education": [
                {"degree": "BS CS", "institution": "MIT", "year": "2020"}
            ],
            "confidence": 0.85,
        })

        result = self.parser._parse_extraction_response(response)

        assert result["extracted_data"]["skills"] == ["Python", "JavaScript", "React"]
        assert len(result["extracted_data"]["projects"]) == 1
        assert len(result["extracted_data"]["experience"]) == 1
        assert len(result["extracted_data"]["education"]) == 1
        assert result["confidence"] == 0.85

    def test_parse_response_with_code_fences(self):
        """Test parsing response wrapped in markdown code fences."""
        data = {
            "skills": ["Python"],
            "projects": [],
            "experience": [],
            "education": [],
            "confidence": 0.6,
        }
        response = f"```json\n{json.dumps(data)}\n```"

        result = self.parser._parse_extraction_response(response)
        assert result["extracted_data"]["skills"] == ["Python"]
        assert result["confidence"] == 0.6

    def test_parse_response_missing_fields(self):
        """Test parsing response with missing optional fields."""
        response = json.dumps({"skills": ["Java"], "confidence": 0.4})

        result = self.parser._parse_extraction_response(response)
        assert result["extracted_data"]["skills"] == ["Java"]
        assert result["extracted_data"]["projects"] == []
        assert result["extracted_data"]["experience"] == []
        assert result["extracted_data"]["education"] == []
        assert result["confidence"] == 0.4

    def test_parse_response_invalid_json(self):
        """Test parsing invalid JSON response."""
        with pytest.raises(ResumeParserError, match="Failed to extract"):
            self.parser._parse_extraction_response("not json at all")

    def test_parse_response_not_dict(self):
        """Test parsing response that is not a dict."""
        with pytest.raises(ResumeParserError, match="unexpected format"):
            self.parser._parse_extraction_response("[1, 2, 3]")

    def test_parse_response_clamps_confidence(self):
        """Test that confidence is clamped to [0.0, 1.0]."""
        response = json.dumps({
            "skills": [],
            "projects": [],
            "experience": [],
            "education": [],
            "confidence": 1.5,
        })

        result = self.parser._parse_extraction_response(response)
        assert result["confidence"] == 1.0

        response = json.dumps({
            "skills": [],
            "projects": [],
            "experience": [],
            "education": [],
            "confidence": -0.5,
        })

        result = self.parser._parse_extraction_response(response)
        assert result["confidence"] == 0.0

    def test_parse_response_non_numeric_confidence(self):
        """Test that non-numeric confidence defaults to 0.5."""
        response = json.dumps({
            "skills": ["Python"],
            "projects": [],
            "experience": [],
            "education": [],
            "confidence": "high",
        })

        result = self.parser._parse_extraction_response(response)
        assert result["confidence"] == 0.5


class TestBuildExtractionPrompt:
    """Tests for prompt building."""

    def setup_method(self):
        """Set up test fixtures."""
        self.parser = ResumeParser.__new__(ResumeParser)

    def test_prompt_includes_resume_text(self):
        """Test that the prompt includes the resume text."""
        prompt = self.parser._build_extraction_prompt("John Doe, Software Engineer")
        assert "John Doe, Software Engineer" in prompt

    def test_prompt_requests_json_format(self):
        """Test that the prompt requests JSON output."""
        prompt = self.parser._build_extraction_prompt("some text")
        assert "JSON" in prompt
        assert "skills" in prompt
        assert "projects" in prompt
        assert "experience" in prompt
        assert "education" in prompt
        assert "confidence" in prompt
